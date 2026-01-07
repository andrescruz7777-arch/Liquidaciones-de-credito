import streamlit as st
import pandas as pd
from datetime import date, timedelta
from decimal import Decimal, ROUND_HALF_UP
from docx import Document
import io
import zipfile
import datetime as dt
from pathlib import Path

# Word helpers
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE_DIR = Path(__file__).parent

# ============================================
#   1. BUSCAR PLANTILLA .DOCX (ROBUSTO)
# ============================================

def obtener_ruta_plantilla() -> str:
    docx_files = list(BASE_DIR.glob("*.docx"))

    st.sidebar.markdown("### 📂 Archivos .docx encontrados:")
    if not docx_files:
        st.sidebar.write("❌ No hay .docx en la raíz")
        st.error(
            "No encontré NINGÚN archivo .docx en la raíz del repo.\n\n"
            "Verifica que subiste la plantilla."
        )
        st.stop()

    for p in docx_files:
        st.sidebar.write("•", p.name)

    plantilla = docx_files[0]
    st.sidebar.markdown(f"**✔ Usando plantilla:** {plantilla.name}")
    return str(plantilla)


# ============================================
#   2. UTILIDADES DE TEXTO (NÚMEROS A LETRAS)
# ============================================

UNIDADES = (
    "cero", "uno", "dos", "tres", "cuatro", "cinco", "seis",
    "siete", "ocho", "nueve", "diez", "once", "doce", "trece",
    "catorce", "quince", "dieciséis", "diecisiete", "dieciocho",
    "diecinueve", "veinte"
)

DECENAS = (
    "cero", "diez", "veinte", "treinta", "cuarenta", "cincuenta",
    "sesenta", "setenta", "ochenta", "noventa"
)

CENTENAS = (
    "cero", "cien", "doscientos", "trescientos", "cuatrocientos",
    "quinientos", "seiscientos", "setecientos",
    "ochocientos", "novecientos"
)

def numero_a_letras_menor_1000(n: int) -> str:
    n = int(n)
    if n < 21:
        return UNIDADES[n]
    if n < 100:
        d, u = divmod(n, 10)
        if u == 0:
            return DECENAS[d]
        if d == 2:
            return "veinti" + UNIDADES[u]
        return DECENAS[d] + " y " + UNIDADES[u]
    c, r = divmod(n, 100)
    if n == 100:
        return "cien"
    pref = CENTENAS[c]
    if r == 0:
        return pref
    return pref + " " + numero_a_letras_menor_1000(r)

def numero_a_letras_centavos(n: int) -> str:
    return numero_a_letras_menor_1000(n)

def numero_a_letras_pesos(valor: float) -> str:
    v = Decimal(str(valor)).quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)
    entero = int(v)
    centavos = int((v - Decimal(entero)) * 100)

    millones, resto = divmod(entero, 1_000_000)
    miles, unidades = divmod(resto, 1_000)

    partes = []
    if millones > 0:
        partes.append("un millón" if millones == 1 else numero_a_letras_menor_1000(millones) + " millones")
    if miles > 0:
        partes.append("mil" if miles == 1 else numero_a_letras_menor_1000(miles) + " mil")
    if unidades > 0 or entero == 0:
        partes.append(numero_a_letras_menor_1000(unidades))

    texto_entero = " ".join(partes)

    if centavos == 0:
        return f"{texto_entero} pesos"
    else:
        return f"{texto_entero} pesos con {numero_a_letras_centavos(centavos)} centavos"


# ============================================
#   3. LECTURA Y NORMALIZACIÓN DE USURA
# ============================================

def cargar_usura(path: str):
    df = pd.read_excel(path)

    if "Fecha desde" in df.columns:
        col_f = "Fecha desde"
    elif "DESDE" in df.columns:
        col_f = "DESDE"
    else:
        st.error("No encontré columna de fecha en TASAS_DE_USURA.xlsx")
        st.stop()

    if "Tasa EA" in df.columns:
        col_t = "Tasa EA"
    elif "TASA DE USURA" in df.columns:
        col_t = "TASA DE USURA"
    else:
        st.error("No encontré columna de tasa en TASAS_DE_USURA.xlsx")
        st.stop()

    meses = {
        "Ene": "Jan", "Feb": "Feb", "Mar": "Mar", "Abr": "Apr",
        "May": "May", "Jun": "Jun", "Jul": "Jul", "Ago": "Aug",
        "Sep": "Sep", "Set": "Sep", "Oct": "Oct", "Nov": "Nov", "Dic": "Dec"
    }

    def parse_fecha(val):
        if isinstance(val, (dt.date, dt.datetime)):
            return val.date() if isinstance(val, dt.datetime) else val
        s = str(val)
        for es, en in meses.items():
            s = s.replace(es, en)
        return pd.to_datetime(s, dayfirst=True).date()

    df["fecha_desde"] = df[col_f].apply(parse_fecha)
    df["tasa_ea"] = df[col_t].astype(float)
    df = df[["fecha_desde", "tasa_ea"]].sort_values("fecha_desde").reset_index(drop=True)
    return df

def obtener_tasa_ea(df_usura, fecha):
    filtro = df_usura[df_usura["fecha_desde"] <= fecha]
    if filtro.empty:
        st.error(f"No hay tasa de usura para la fecha {fecha}")
        st.stop()
    return Decimal(str(filtro.iloc[-1]["tasa_ea"]))


# ============================================
#   4. MOTOR DE LIQUIDACIÓN
# ============================================

def liquidar_obligacion(fila, df_usura, fecha_liq):

    capital = Decimal(str(fila["CAPITAL"]))

    fecha_venc = pd.to_datetime(fila["FECHA VENCIMIENTO PAGARÉ"]).date()
    fecha_intereses = fecha_venc + timedelta(days=1)

    fecha_actual = fecha_intereses
    interes_acum = Decimal("0")
    filas = []

    while fecha_actual <= fecha_liq:

        fin_mes = (fecha_actual.replace(day=1) + timedelta(days=32)).replace(day=1) - timedelta(days=1)
        fecha_hasta = min(fin_mes, fecha_liq)

        dias = (fecha_hasta - fecha_actual).days + 1
        tasa_ea = obtener_tasa_ea(df_usura, fecha_actual)

        factor = ((Decimal("1") + tasa_ea) ** (Decimal("1") / Decimal("365"))) - Decimal("1")

        interes_periodo = (capital * factor * dias).quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)
        interes_acum += interes_periodo

        filas.append({
            "fecha_desde": fecha_actual,
            "fecha_hasta": fecha_hasta,
            "tasa_ea": float(tasa_ea),
            "factor_dia": float(factor),
            "dias": dias,
            "interes_periodo": float(interes_periodo),
            "interes_acumulado": float(interes_acum)
        })

        fecha_actual = fecha_hasta + timedelta(days=1)

    df_detalle = pd.DataFrame(filas)

    resumen = {
        "nombre": fila["NOMBRE"],
        "cedula": fila["CEDULA"],
        "pagaré": fila["No. PAGARÉ"],
        "juzgado": fila["JUZGADO"],
        "correo_juzgado": fila["CORREO JUZGADO"],
        "radicado": fila["RADICADO"],
        "capital": float(capital),
        "total_mora": float(interes_acum),
        "saldo_total": float(capital + interes_acum),
        "fecha_intereses": fecha_intereses,
        "fecha_liquidacion": fecha_liq
    }

    return df_detalle, resumen


# ============================================
#   5. WORD: REEMPLAZO SIN ROMPER FORMATO
# ============================================

def _copiar_formato(origen_run, destino_run):
    destino_run.bold = origen_run.bold
    destino_run.italic = origen_run.italic
    destino_run.underline = origen_run.underline
    destino_run.font.name = origen_run.font.name
    destino_run.font.size = origen_run.font.size
    if origen_run.font.color and origen_run.font.color.rgb:
        destino_run.font.color.rgb = origen_run.font.color.rgb

def _replace_placeholder_en_parrafo(paragraph, placeholder: str, value: str):
    full_text = "".join(r.text for r in paragraph.runs)
    if placeholder not in full_text:
        return

    start = full_text.find(placeholder)
    end = start + len(placeholder)

    idx = 0
    runs = paragraph.runs
    first_run_idx = last_run_idx = None
    start_in_run = end_in_run = None

    for i, r in enumerate(runs):
        run_len = len(r.text)
        run_start = idx
        run_end = idx + run_len

        if first_run_idx is None and start >= run_start and start <= run_end:
            first_run_idx = i
            start_in_run = start - run_start

        if first_run_idx is not None and end >= run_start and end <= run_end:
            last_run_idx = i
            end_in_run = end - run_start
            break

        idx += run_len

    if first_run_idx is None or last_run_idx is None:
        return

    # Placeholder en un solo run
    if first_run_idx == last_run_idx:
        r = runs[first_run_idx]
        r.text = r.text[:start_in_run] + value + r.text[end_in_run:]
        return

    # Placeholder atraviesa varios runs
    first_run = runs[first_run_idx]
    last_run = runs[last_run_idx]

    prefix = first_run.text[:start_in_run]
    suffix = last_run.text[end_in_run:]

    first_run.text = prefix
    for j in range(first_run_idx + 1, last_run_idx):
        runs[j].text = ""
    last_run.text = suffix

    new_run = paragraph.add_run(value)
    _copiar_formato(first_run, new_run)
    first_run._element.addnext(new_run._element)

def reemplazar(doc, placeholder, valor):
    for p in doc.paragraphs:
        _replace_placeholder_en_parrafo(p, placeholder, valor)

    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _replace_placeholder_en_parrafo(p, placeholder, valor)

def limpiar_parrafos_vacios_final(doc, max_limpiar=80):
    # Limpieza agresiva de párrafos vacíos al final (evita páginas raras)
    count = 0
    while doc.paragraphs and count < max_limpiar:
        p = doc.paragraphs[-1]
        if (p.text or "").strip() == "":
            p._element.getparent().remove(p._element)
            count += 1
        else:
            break


# ============================================
#   6. TABLA: FORMATO
# ============================================

def aplicar_estilo_tabla(tabla):
    estilos_preferidos = [
        "Table Grid",
        "Grid Table 4 Accent 1",
        "Light Shading Accent 1",
        "Light List Accent 1",
    ]
    for s in estilos_preferidos:
        try:
            tabla.style = s
            break
        except Exception:
            continue

    # Encabezado: negrilla y centrado
    for cell in tabla.rows[0].cells:
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True

    # Cuerpo: columnas centradas y valores a la derecha
    for row in tabla.rows[1:]:
        for i, cell in enumerate(row.cells):
            for p in cell.paragraphs:
                if i in (0, 1, 2, 3, 4):
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT


# ============================================
#   7. GENERADOR DE MEMORIAL
# ============================================

def generar_memorial(resumen, df_detalle):
    ruta = obtener_ruta_plantilla()
    doc = Document(ruta)

    reemplazar(doc, "{{JUZGADO}}", str(resumen["juzgado"]))
    reemplazar(doc, "{{CORREO_JUZGADO}}", str(resumen["correo_juzgado"]))
    reemplazar(doc, "{{RADICADO}}", str(resumen["radicado"]))
    reemplazar(doc, "{{NOMBRE}}", str(resumen["nombre"]))
    reemplazar(doc, "{{CEDULA}}", str(resumen["cedula"]))
    reemplazar(doc, "{{PAGARE}}", str(resumen["pagaré"]))

    reemplazar(doc, "{{FECHA_INTERESES}}", resumen["fecha_intereses"].strftime("%d/%m/%Y"))
    reemplazar(doc, "{{FECHA_LIQUIDACION}}", resumen["fecha_liquidacion"].strftime("%d/%m/%Y"))

    reemplazar(doc, "{{CAPITAL}}", f"${resumen['capital']:,.2f}")
    reemplazar(doc, "{{TOTAL_MORA}}", f"${resumen['total_mora']:,.2f}")
    reemplazar(doc, "{{SALDO_TOTAL}}", f"${resumen['saldo_total']:,.2f}")

    # ✅ MAYÚSCULAS
    letras = numero_a_letras_pesos(resumen["saldo_total"]).upper()
    reemplazar(doc, "{{VALOR_LETRAS}}", letras)
    reemplazar(doc, "{{VALOR_NUM}}", f"${resumen['saldo_total']:,.2f}")

    # ✅ Limpieza final (evita acumulación de enters que generan páginas vacías)
    limpiar_parrafos_vacios_final(doc)

    # ✅ NO forzamos salto de página: Word baja la tabla si no cabe (sin páginas en blanco)
    doc.add_paragraph("")  # separador suave

    tabla = doc.add_table(rows=1, cols=7)
    h = tabla.rows[0].cells
    h[0].text, h[1].text, h[2].text, h[3].text, h[4].text, h[5].text, h[6].text = (
        "Desde", "Hasta", "EA", "Factor día", "Días", "Interés período", "Acumulado"
    )

    for _, r in df_detalle.iterrows():
        row = tabla.add_row().cells
        row[0].text = r["fecha_desde"].strftime("%d/%m/%Y")
        row[1].text = r["fecha_hasta"].strftime("%d/%m/%Y")
        row[2].text = f"{r['tasa_ea']*100:.2f}%"
        row[3].text = f"{r['factor_dia']*100:.5f}%"
        row[4].text = str(r["dias"])
        row[5].text = f"${r['interes_periodo']:,.2f}"
        row[6].text = f"${r['interes_acumulado']:,.2f}"

    aplicar_estilo_tabla(tabla)

    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()


# ============================================
#   8. INTERFAZ STREAMLIT
# ============================================

st.title("💼 Liquidador Judicial Masivo – Banco GNB Sudameris")

st.subheader("1️⃣ Cargar base de obligaciones")
archivo_base = st.file_uploader("Sube el archivo Excel", type=["xlsx"])

if archivo_base:
    df_base = pd.read_excel(archivo_base)
    st.success(f"Base cargada: {len(df_base)} registros")

    cols = [
        "NOMBRE", "CEDULA", "JUZGADO", "CORREO JUZGADO",
        "RADICADO", "FECHA VENCIMIENTO PAGARÉ",
        "CAPITAL", "No. PAGARÉ"
    ]
    faltan = [c for c in cols if c not in df_base.columns]
    if faltan:
        st.error(f"Faltan columnas: {faltan}")
        st.stop()

    st.subheader("2️⃣ Fecha de liquidación")
    fecha_liq = st.date_input("Seleccione fecha", value=date.today())

    st.subheader("3️⃣ Cargar tasas de usura")
    df_usura = cargar_usura("TASAS_DE_USURA.xlsx")
    st.success("Usura cargada correctamente")

    st.subheader("4️⃣ Selecciona una obligación para previsualizar")
    lista_pagare = df_base["No. PAGARÉ"].astype(str).tolist()
    pag = st.selectbox("Obligación:", lista_pagare)

    fila = df_base[df_base["No. PAGARÉ"].astype(str) == pag].iloc[0]
    df_det, resumen = liquidar_obligacion(fila, df_usura, fecha_liq)

    st.markdown("### 🔍 Resumen")
    st.json({
        "Cliente": resumen["nombre"],
        "Identificación": resumen["cedula"],
        "Pagaré": resumen["pagaré"],
        "Fecha intereses": resumen["fecha_intereses"].strftime("%d/%m/%Y"),
        "Fecha liquidación": resumen["fecha_liquidacion"].strftime("%d/%m/%Y"),
        "Capital": f"${resumen['capital']:,.2f}",
        "Total mora": f"${resumen['total_mora']:,.2f}",
        "Saldo total": f"${resumen['saldo_total']:,.2f}",
        "En letras (MAYÚSCULAS)": numero_a_letras_pesos(resumen["saldo_total"]).upper()
    })

    st.markdown("### 📊 Detalle por períodos")
    df_vista = df_det.copy()
    df_vista["tasa_ea"] = df_vista["tasa_ea"].map(lambda x: f"{x*100:.2f}%")
    df_vista["factor_dia"] = df_vista["factor_dia"].map(lambda x: f"{x*100:.5f}%")
    df_vista["interes_periodo"] = df_vista["interes_periodo"].map(lambda x: f"${x:,.2f}")
    df_vista["interes_acumulado"] = df_vista["interes_acumulado"].map(lambda x: f"${x:,.2f}")
    st.dataframe(df_vista)

    st.subheader("5️⃣ Generar memorial individual")
    if st.button("Generar memorial"):
        docx_bytes = generar_memorial(resumen, df_det)
        st.download_button(
            "📄 Descargar memorial",
            docx_bytes,
            file_name=f"MEMORIAL_{resumen['pagaré']}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    st.subheader("6️⃣ Generar memoriales masivos")
    if st.button("Generar ZIP masivo"):
        buffer = io.BytesIO()
        with zipfile.ZipFile(buffer, "w") as z:
            for _, f in df_base.iterrows():
                det, res = liquidar_obligacion(f, df_usura, fecha_liq)
                archivo = generar_memorial(res, det)
                nombre = f"MEMORIAL_{res['pagaré']}.docx"
                z.writestr(nombre, archivo)

        buffer.seek(0)
        st.download_button(
            "📦 Descargar ZIP",
            buffer.getvalue(),
            file_name="MEMORIALES_GNB.zip",
            mime="application/zip"
        )
else:
    st.info("Sube primero la base .xlsx para comenzar.")
